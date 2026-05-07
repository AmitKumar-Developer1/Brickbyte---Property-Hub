from __future__ import annotations

import os
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "BrickByte_Project_Explanation_and_Presentation_Guide.docx"

ACCENT = RGBColor(212, 117, 91)
DARK = RGBColor(28, 27, 26)
MUTED = RGBColor(92, 88, 84)
LIGHT_FILL = "FAF8F4"
TABLE_HEADER = "F5F1E8"


EXCLUDED_DIRS = {
    "node_modules",
    "dist",
    ".git",
    ".npm-cache",
    ".agent",
    ".agents",
    "mongo-local",
    "mongo-local-package",
    "Real-Estate-Website-main",
}

GENERATED_EXTS = {".log", ".bak"}


def rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def iter_project_files() -> list[Path]:
    files: list[Path] = []
    for dirpath, dirnames, filenames in os.walk(ROOT):
        p = Path(dirpath)
        dirnames[:] = [d for d in dirnames if d not in EXCLUDED_DIRS]
        if "docs" in p.parts or "tools" in p.parts:
            continue
        for name in filenames:
            fp = p / name
            if fp.suffix in GENERATED_EXTS:
                continue
            files.append(fp)
    return sorted(files, key=lambda x: rel(x).lower())


def classify_language(path: Path) -> str:
    ext = path.suffix.lower()
    mapping = {
        ".tsx": "TypeScript + JSX",
        ".ts": "TypeScript",
        ".jsx": "JavaScript + JSX",
        ".js": "JavaScript",
        ".mjs": "JavaScript module",
        ".css": "CSS",
        ".html": "HTML",
        ".json": "JSON",
        ".md": "Markdown",
        ".yml": "YAML",
        ".yaml": "YAML",
        ".sh": "Shell script",
        ".svg": "SVG vector",
        ".png": "PNG image",
        ".jpg": "JPEG image",
        ".jpeg": "JPEG image",
        ".ico": "Icon",
        ".webmanifest": "Web manifest",
        ".xml": "XML",
        ".txt": "Text",
        ".example": "Environment template",
    }
    if path.name == "LICENSE":
        return "License text"
    if path.name.startswith(".env"):
        return "Environment file"
    if path.name == ".gitignore":
        return "Git ignore rules"
    return mapping.get(ext, ext[1:].upper() if ext else "Plain text")


def read_text(path: Path, limit: int = 200_000) -> str:
    try:
        if path.stat().st_size > limit:
            return ""
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


def extract_symbols(path: Path) -> str:
    text = read_text(path)
    if not text:
        return ""
    symbols: list[str] = []
    for pat in [
        r"export\s+default\s+function\s+([A-Za-z0-9_]+)",
        r"export\s+function\s+([A-Za-z0-9_]+)",
        r"export\s+async\s+function\s+([A-Za-z0-9_]+)",
        r"export\s+const\s+([A-Za-z0-9_]+)",
        r"const\s+([A-Z][A-Za-z0-9_]+)\s*[:=]",
        r"class\s+([A-Za-z0-9_]+)",
        r"interface\s+([A-Za-z0-9_]+)",
    ]:
        symbols.extend(re.findall(pat, text))
    seen = []
    for s in symbols:
        if s not in seen:
            seen.append(s)
    return ", ".join(seen[:10])


def extract_api_calls(path: Path) -> str:
    text = read_text(path)
    if not text:
        return ""
    matches = re.findall(r"(?:apiClient|axios)\.(?:get|post|put|delete)\(([`'\"])(.+?)\1", text)
    paths = [m[1] for m in matches]
    matches2 = re.findall(r"apiClient\.(?:get|post|put|delete)\((`.+?`)", text)
    paths.extend(matches2)
    seen = []
    for p in paths:
        p = p.replace("${id}", ":id").replace("${queryParams}", "query")
        if p not in seen:
            seen.append(p)
    return ", ".join(seen[:8])


def extract_routes(path: Path) -> list[tuple[str, str, str]]:
    text = read_text(path)
    routes = []
    if not text:
        return routes
    for m in re.finditer(
        r"(?:router|userrouter|propertyrouter|newsrouter)\.(get|post|put|delete)\((['\"])(.*?)\2\s*,\s*([^)]+)\)",
        text,
        flags=re.S,
    ):
        method = m.group(1).upper()
        route_path = m.group(3)
        handlers = " ".join(m.group(4).replace("\n", " ").split())
        handlers = handlers[:120]
        routes.append((method, route_path, handlers))
    return routes


def purpose_for(path: Path) -> str:
    r = rel(path)
    name = path.name
    parent = path.parent.name
    text = read_text(path, limit=80_000)

    explicit = {
        "frontend/src/App.tsx": "Root customer app: BrowserRouter, AuthProvider, route table, lazy page loading, page transitions, global toaster.",
        "frontend/src/main.tsx": "Customer frontend entry point: finds #root in index.html and renders React StrictMode + App.",
        "frontend/src/services/api.ts": "Central Axios client for customer app; sets base URL, auth token interceptor, AI API key headers, and grouped API helper objects.",
        "frontend/src/contexts/AuthContext.tsx": "Customer auth state provider; reads/writes JWT and user data in localStorage and exposes login/register/logout.",
        "frontend/src/hooks/useSEO.ts": "Small hook that changes document title and meta description per route.",
        "frontend/src/components/common/StructuredData.tsx": "Injects JSON-LD schema.org metadata for website, organization, property, and AI Hub pages.",
        "backend/server.js": "Express server bootstrap: environment loading, security middleware, CORS, MongoDB startup, route mounting, health/status endpoints, shutdown handlers.",
        "backend/serverweb.js": "Generates the backend root/status HTML page shown at /.",
        "backend/email.js": "Standalone or legacy email helper/template code used around mail delivery setup.",
        "backend/config/mongodb.js": "Mongoose connection manager with retry, connection listeners, health checks, sanitized URI output, and close hooks.",
        "backend/config/imagekit.js": "ImageKit SDK configuration and required-key validation for image uploads.",
        "backend/config/nodemailer.js": "SMTP transporter configuration used for transactional emails.",
        "backend/config/config.js": "Central config export for API/service keys and environment-dependent values.",
        "admin/src/App.jsx": "Root admin dashboard app: protected route layout, sidebar state, page transitions, error boundary, and Sonner toaster.",
        "admin/src/main.jsx": "Admin frontend entry point: renders React app into index.html.",
        "admin/src/services/apiClient.js": "Admin Axios client; injects admin JWT from localStorage and clears auth on 401 responses.",
        "admin/src/contexts/AuthContext.jsx": "Admin auth provider; decodes JWT, tracks admin login state, and handles logout.",
        "shared/utils/formatPrice.ts": "Reusable INR price formatter intended for shared/frontend logic.",
        "shared/utils/index.ts": "Shared utility barrel export.",
    }
    if r in explicit:
        return explicit[r]

    if r.startswith(".github/"):
        if "ISSUE_TEMPLATE" in r:
            return "GitHub issue template for project maintenance workflow."
        if name == "PULL_REQUEST_TEMPLATE.md":
            return "Pull request checklist/template for contributors."
        if name == "CODEOWNERS":
            return "GitHub CODEOWNERS ownership mapping."

    if parent == "routes" and path.suffix == ".js":
        routes = extract_routes(path)
        route_desc = ", ".join([f"{m} {p}" for m, p, _ in routes[:6]])
        return f"Express router that maps HTTP endpoints to controller functions. Routes: {route_desc}."

    if parent == "controller" and path.suffix == ".js":
        syms = extract_symbols(path)
        return f"Backend controller module containing request handlers. Key exported handlers: {syms or 'route handlers'}."

    if parent == "models" and path.suffix == ".js":
        if "new mongoose.Schema" in text:
            return "Mongoose data model/schema; defines collection fields, validation, timestamps, and indexes used by controllers."
        return "MongoDB/Mongoose model helper."

    if parent == "services" and path.suffix == ".js":
        if "Firecrawl" in text:
            return "Service wrapper around Firecrawl search/scrape APIs; builds real estate queries, extracts structured listing data, filters results, and handles retries/circuit breakers."
        if "ModelClient" in text or "GitHub Models" in text:
            return "AI service wrapper for GitHub Models/Azure inference; validates keys, calls GPT-4.1 models, builds analysis prompts, and falls back with circuit breakers."
        if "nodemailer" in text or "send" in text.lower():
            return "Email service class/templates for password reset, welcome, appointment, and admin/user notification emails."
        return "Backend service abstraction used by controllers."

    if parent == "middleware" and path.suffix == ".js":
        return f"Express middleware for {name.replace('.js', '').replace('Middleware', '').lower()} concerns in the API request pipeline."

    if parent == "utils" and path.suffix in {".js", ".ts"}:
        return f"Utility module for {name.replace('.js', '').replace('.ts', '')}; reused by controllers/services to keep repeated logic out of route handlers."

    if r.startswith("frontend/src/pages/"):
        page = name.replace("Page.tsx", "").replace(".tsx", "")
        return f"Customer-facing React page for the {page} route; composes layout, SEO, API calls/state, and section components."

    if r.startswith("frontend/src/components/ai-hub/"):
        return "AI Hub React component: search form, API key modal, scraped result cards, analysis sidebar, trend display, or AI call-to-action section."

    if r.startswith("frontend/src/components/auth/"):
        return "Authentication UI component for customer sign-in/sign-up headers, form state, validation, and social-login button layout."

    if r.startswith("frontend/src/components/common/"):
        return "Shared customer UI component used across pages: navigation, footer, loading/transition behavior, scroll handling, or SEO metadata."

    if r.startswith("frontend/src/components/home/"):
        return "Homepage section component; displays landing content such as hero, stats, process, trust signals, curated listings, AI section, testimonials, or CTA."

    if r.startswith("frontend/src/components/about/"):
        return "About page section component; explains heritage, AI, values, stats, hero, or CTA content."

    if r.startswith("frontend/src/components/contact/"):
        return "Contact page section component; handles form, contact cards, FAQ, map area, newsletter banner, or other contact options."

    if r.startswith("frontend/src/components/properties/"):
        return "Property listing UI component for filter controls, listing header, grid/list rendering, or individual property cards."

    if r.startswith("frontend/src/components/property-details/"):
        return "Property detail page section component for breadcrumbs, hero image, facts, amenities, location, and viewing schedule form."

    if r.startswith("frontend/src/components/ui/"):
        return "Reusable design-system primitive, mostly Radix/shadcn-style wrapper components with Tailwind classes."

    if r.startswith("frontend/src/components/figma/"):
        return "Figma-import helper component for safe image rendering with fallback behavior."

    if r.startswith("frontend/src/styles/"):
        return "Frontend stylesheet for fonts, Tailwind layer imports, theme tokens, or global CSS."

    if r.startswith("frontend/src/assets/"):
        return "Hashed image asset bundled by Vite/Figma import pipeline; used as static visual media in React components."

    if r.startswith("frontend/src/images/"):
        return "Named page image asset used by hero cards, about/home sections, listing examples, or contact visuals."

    if r.startswith("frontend/public/"):
        return "Public static asset served directly by Vite: favicon, logo, manifest, sitemap, robots file, or Open Graph image."

    if r.startswith("frontend/Assets/"):
        return "Brand asset source copy for logos, favicons, and Open Graph graphics."

    if r.startswith("admin/src/pages/"):
        page = name.replace(".jsx", "")
        return f"Admin dashboard page for {page}: fetches backend admin APIs, renders tables/cards/forms, and handles management actions."

    if r.startswith("admin/src/components/"):
        return "Admin UI component for layout, authentication guard/login, modals, bulk action controls, status badges, or error fallback."

    if r.startswith("admin/src/"):
        return "Admin app source file for configuration, constants, hook, utility, styling, or React entry/layout behavior."

    if r.startswith("Image/"):
        return "README/documentation screenshot used to show platform preview and AI Hub views."

    if name == "package.json":
        return "NPM manifest listing scripts, module type, dependencies, dev dependencies, author/license metadata."
    if name == "package-lock.json":
        return "Locked dependency tree for reproducible npm installs."
    if name.startswith(".env"):
        return "Environment configuration/template; controls URLs, feature flags, secrets, service keys, and deployment settings."
    if name in {"Dockerfile", "docker-compose.yml"}:
        return "Container/runtime configuration for local or hosted deployment."
    if name in {"render.yaml", "vercel.json"}:
        return "Deployment platform configuration."
    if name.startswith("vite.config"):
        return "Vite build/dev-server configuration."
    if name.startswith("tailwind.config") or name.startswith("postcss.config"):
        return "CSS tooling configuration."
    if name.startswith("tsconfig"):
        return "TypeScript compiler configuration."
    if name == "index.html":
        return "HTML shell containing metadata and the root div where the React app mounts."
    if path.suffix.lower() == ".md":
        return "Markdown documentation for setup, contribution, security, architecture notes, or project overview."
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".svg", ".ico"}:
        return "Static visual asset."
    return "Project file used by the BrickByte monorepo."


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.color.rgb = ACCENT
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = DARK
            run.font.size = Pt(13)
        else:
            run.font.color.rgb = DARK
            run.font.size = Pt(11)


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = "Arial"
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(10)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(10)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.05)
    for line in text.strip("\n").splitlines():
        run = p.add_run(line.rstrip() + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(45, 45, 45)
    p.paragraph_format.space_after = Pt(8)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=font_size, color=DARK)
        set_cell_fill(hdr[i], TABLE_HEADER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Subtitle"].font.name = "Arial"
    styles["Subtitle"].font.size = Pt(12)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "BrickByte Project Guide"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if header.runs:
            header.runs[0].font.name = "Arial"
            header.runs[0].font.size = Pt(8)
            header.runs[0].font.color.rgb = MUTED
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.text = "Generated project explanation"
        if footer.runs:
            footer.runs[0].font.name = "Arial"
            footer.runs[0].font.size = Pt(8)
            footer.runs[0].font.color.rgb = MUTED


def summarize_inventory(files: list[Path]) -> dict[str, list[Path]]:
    groups: dict[str, list[Path]] = defaultdict(list)
    for f in files:
        r = rel(f)
        if r.startswith("frontend/src/components/ui/"):
            groups["frontend ui primitives"].append(f)
        elif r.startswith("frontend/src/components/"):
            groups["frontend components"].append(f)
        elif r.startswith("frontend/src/pages/"):
            groups["frontend pages"].append(f)
        elif r.startswith("frontend/src/"):
            groups["frontend source core"].append(f)
        elif r.startswith("frontend/public/") or r.startswith("frontend/Assets/"):
            groups["frontend public assets"].append(f)
        elif r.startswith("admin/src/pages/"):
            groups["admin pages"].append(f)
        elif r.startswith("admin/src/components/"):
            groups["admin components"].append(f)
        elif r.startswith("admin/src/"):
            groups["admin source core"].append(f)
        elif r.startswith("backend/"):
            if "/controller/" in r:
                groups["backend controllers"].append(f)
            elif "/routes/" in r:
                groups["backend routes"].append(f)
            elif "/models/" in r:
                groups["backend models"].append(f)
            elif "/services/" in r:
                groups["backend services"].append(f)
            elif "/middleware/" in r:
                groups["backend middleware"].append(f)
            elif "/utils/" in r:
                groups["backend utils"].append(f)
            elif "/config/" in r:
                groups["backend config"].append(f)
            else:
                groups["backend root/deploy"].append(f)
        elif r.startswith("shared/"):
            groups["shared utilities"].append(f)
        elif r.startswith("Image/"):
            groups["documentation images"].append(f)
        elif r.startswith(".github/"):
            groups["github workflow docs"].append(f)
        elif r.startswith("admin/"):
            groups["admin root/deploy"].append(f)
        elif r.startswith("frontend/"):
            groups["frontend root/deploy"].append(f)
        else:
            groups["repo root docs"].append(f)
    return dict(groups)


def route_rows() -> list[list[str]]:
    mount_prefix = {
        "backend/routes/productRoutes.js": "/api/products",
        "backend/routes/userRoutes.js": "/api/users",
        "backend/routes/formRoutes.js": "/api/forms",
        "backend/routes/newsRoutes.js": "/api/news",
        "backend/routes/appointmentRoutes.js": "/api/appointments",
        "backend/routes/adminRoutes.js": "/api/admin",
        "backend/routes/propertyRoutes.js": "/api",
        "backend/routes/healthRoutes.js": "/health",
    }
    rows = []
    for relp, prefix in mount_prefix.items():
        path = ROOT / relp
        for method, sub, handlers in extract_routes(path):
            rows.append([method, prefix + (sub if sub.startswith("/") else "/" + sub), handlers])
    return rows


def frontend_routes() -> list[list[str]]:
    return [
        ["/", "HomePage", "Landing page with hero, stats, curated listings, AI section, process, trust, testimonials, CTA."],
        ["/properties", "PropertiesPage", "Fetches active listings from backend, filters/sorts locally, renders grid/list cards."],
        ["/property/:id", "PropertyDetailsPage", "Fetches a single property, renders gallery/details/amenities/location, supports appointment booking."],
        ["/ai-hub", "AIPropertyHubPage", "Feature-flagged AI Hub; runs full search locally when VITE_ENABLE_AI_HUB=true."],
        ["/about", "AboutUsPage", "Company/about storytelling sections and SEO metadata."],
        ["/contact", "ContactPage", "Contact sections and form submission to backend."],
        ["/signin", "SignInPage", "Customer login route using AuthContext/userAPI."],
        ["/signup", "SignUpPage", "Customer registration route using AuthContext/userAPI."],
        ["/forgot-password", "ForgotPasswordPage", "Requests reset email via /api/users/forgot."],
        ["/reset/:token", "ResetPasswordPage", "Submits new password via tokenized reset route."],
        ["/verify-email/:token", "VerifyEmailPage", "Calls email verification endpoint and stores returned token if successful."],
        ["/add-property", "AddPropertyPage", "Authenticated user listing form; uploads images using FormData."],
        ["/my-listings", "MyListingsPage", "Authenticated user's own submitted listings and delete workflow."],
    ]


def admin_routes() -> list[list[str]]:
    return [
        ["/login", "Login", "Admin login form calling /api/users/admin."],
        ["/dashboard", "Dashboard", "Overview stats, charts, activity, and quick actions."],
        ["/list", "PropertyListings", "Admin-managed active listing table/cards with delete/update actions."],
        ["/add", "Add", "Admin property creation form with image upload."],
        ["/update/:id", "Update", "Admin property edit form using product single/update APIs."],
        ["/appointments", "Appointments", "View, confirm/cancel/complete appointments and add meeting links."],
        ["/pending-listings", "PendingListings", "Review queue for user-submitted listings; approve/reject/bulk actions."],
        ["/users", "Users", "User search/filter/status management and bulk suspend/ban actions."],
        ["/users/:id", "UserDetails", "Detailed user view with status actions."],
        ["/activity-logs", "ActivityLogs", "Admin activity history with export support."],
    ]


def model_rows() -> list[list[str]]:
    return [
        ["User", "name, email, password, resetToken, isEmailVerified, status, suspension/ban fields, lastActive, timestamps", "Customer account identity, auth, email reset/verify, admin moderation status."],
        ["Admin", "email, password, role, lastLogin", "Admin login identity; password is hashed in a pre-save hook."],
        ["Property", "title, location, price, image[], beds, baths, sqft, type, availability, description, amenities[], phone, googleMapLink, status, postedBy, rejectionReason, expiresAt, timestamps", "Main property listing model used by admin listings, public listing pages, and user-submitted approval workflow."],
        ["Appointment", "propertyId, userId, guestInfo, date, time, status, meetingLink/platform, notes, cancelReason, reminderSent, feedback, timestamps", "Viewing bookings from property detail pages and admin appointment management."],
        ["AdminActivityLog", "adminEmail, action, targetType, targetId, targetName, metadata, ipAddress, userAgent, timestamps", "Audit trail for admin moderation actions and bulk operations."],
        ["SearchCache", "cacheKey, data, createdAt TTL", "MongoDB TTL cache for AI search/trends responses."],
        ["Stats", "endpoint, method, timestamp, responseTime, statusCode, timestamps", "API analytics captured by stats middleware for dashboard/reporting."],
        ["Form", "name, email, phone, message, timestamps", "Contact form submissions."],
        ["News", "email", "Newsletter signup emails."],
    ]


def add_cover(doc: Document, file_count: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BrickByte Project Explanation")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Source Code Architecture, File-by-File Guide, Data Flow, API Integration, and Presentation Preparation")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK

    add_para(doc, "")
    rows = [
        ["Project path", str(ROOT)],
        ["Generated", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["Active apps", "frontend (customer React + TypeScript), admin (React + JavaScript), backend (Node/Express/MongoDB)"],
        ["Files analyzed", f"{file_count} active project-owned files, excluding node_modules/dist/logs/local MongoDB binaries and treating Real-Estate-Website-main as a duplicate snapshot."],
        ["Local URL checked earlier", "http://localhost:5173"],
    ]
    add_table(doc, ["Field", "Value"], rows, widths=[1.7, 5.8], font_size=9)
    add_para(
        doc,
        "Important academic note: use this document to understand, customize, and explain the project. If your institute asks about authorship or sources, be honest and give credit where required.",
    )


def add_architecture_sections(doc: Document, files: list[Path]) -> None:
    add_heading(doc, "1. Executive Summary")
    add_para(
        doc,
        "BrickByte is a full-stack real estate platform. The customer frontend lets users browse properties, view details, create accounts, schedule viewings, submit their own listings, and run an AI-assisted property search. The admin dashboard lets administrators add/update/delete properties, review pending user submissions, manage users, view appointments, and audit admin activity. The backend is an Express API connected to MongoDB, ImageKit, email SMTP, Firecrawl, and GitHub Models."
    )
    add_bullets(
        doc,
        [
            "Frontend: React 18 + TypeScript + Vite + Tailwind CSS v4, with React Router v7 and Framer Motion page transitions.",
            "Admin: React 18 + JavaScript + Vite + Tailwind CSS v3, with Chart.js dashboards and protected admin routes.",
            "Backend: Node.js ESM + Express + MongoDB/Mongoose, with JWT auth, Helmet, CORS, rate limiting, request IDs, Winston logging, image uploads, cron jobs, and AI/search services.",
            "AI Hub: user-provided Firecrawl and GitHub Models keys are saved in browser localStorage and sent as request headers; the backend refuses to use server fallback keys for AI calls.",
            "Duplicated folder: Real-Estate-Website-main is a second copy of the same repo. Key file hashes match, so this guide explains the active top-level apps once.",
        ],
    )

    add_heading(doc, "2. Technology and File Types")
    ext_counts = Counter(f.suffix.lower() or "(no extension)" for f in files)
    rows = []
    for ext, count in ext_counts.most_common():
        sample = next((rel(f) for f in files if (f.suffix.lower() or "(no extension)") == ext), "")
        rows.append([ext, str(count), classify_language(Path("x" + ("" if ext == "(no extension)" else ext))), sample])
    add_table(doc, ["Extension", "Count", "Meaning", "Example"], rows, widths=[0.9, 0.7, 1.7, 4.2], font_size=8)

    add_heading(doc, "3. Monorepo Map")
    add_table(
        doc,
        ["Folder", "Role", "What to say in presentation"],
        [
            ["frontend/", "Customer website", "React/TypeScript SPA: routes, pages, design components, Axios service layer, auth context, SEO metadata, AI Hub UI."],
            ["admin/", "Admin dashboard", "Separate React/Vite app for internal property, user, appointment, and audit management."],
            ["backend/", "REST API server", "Express/Mongoose API that owns business logic, auth, uploads, scraping, AI analysis, caching, logging, and cron jobs."],
            ["shared/", "Small shared utility folder", "Contains price-format helper exports. The architecture notes say local copies are used for production deployment separation."],
            ["Image/", "Documentation screenshots", "Images used by README and project previews."],
            [".github/", "Repository collaboration files", "Issue templates, PR template, and code-owner style repository metadata."],
            ["Real-Estate-Website-main/", "Duplicate source snapshot", "A copied copy of the same project; not the active server folder."],
            ["mongo-local/, mongo-local-package/", "Local MongoDB installer/package files", "Large local runtime/download artifacts, not application source code."],
        ],
        widths=[1.7, 2.1, 3.7],
        font_size=8,
    )

    add_heading(doc, "4. System Architecture Flow")
    add_code_block(
        doc,
        """
flowchart LR
  UserBrowser["Customer browser: React + TS"] -->|"Axios /api/*"| Backend["Express API on port 4000"]
  AdminBrowser["Admin browser: React + JS"] -->|"Axios /api/admin + /api/products"| Backend
  Backend -->|"Mongoose queries"| MongoDB[("MongoDB brickbyte database")]
  Backend -->|"Image upload URLs"| ImageKit["ImageKit CDN"]
  Backend -->|"SMTP templates"| Email["Brevo/Gmail SMTP via nodemailer"]
  Backend -->|"Search + extract listings"| Firecrawl["Firecrawl API"]
  Backend -->|"Rank + analyze JSON"| GitHubModels["GitHub Models GPT-4.1 mini/nano"]
  Backend -->|"Structured JSON"| UserBrowser
  Backend -->|"Stats / audit data"| AdminBrowser
        """,
    )

    add_heading(doc, "5. Browser Boot Flow")
    add_code_block(
        doc,
        """
index.html
  -> <div id="root"></div>
  -> /src/main.tsx
      -> createRoot(root).render(<React.StrictMode><App /></React.StrictMode>)
      -> App.tsx
          -> BrowserRouter
          -> ScrollToTop
          -> AuthProvider
          -> Suspense lazy page loader
          -> AnimatedRoutes with React Router paths
          -> Sonner Toaster
        """,
    )
    add_para(
        doc,
        "The HTML file is mostly metadata and a mount point. React owns the visible app after main.tsx renders App.tsx. App.tsx lazy-loads pages so the initial JavaScript bundle can be smaller, and AnimatePresence/PageTransition add route transition animations."
    )

    add_heading(doc, "6. Backend Request Pipeline")
    add_code_block(
        doc,
        """
HTTP request
  -> global rate limiter
  -> helmet security headers
  -> compression
  -> express.json/urlencoded body parsing
  -> requestIdMiddleware
  -> statsMiddleware
  -> mongoSanitize
  -> CORS allowlist
  -> route module mounted in server.js
  -> controller function
  -> model/service/util layer
  -> JSON response
  -> error handler or 404 fallback when needed
        """,
    )

    add_heading(doc, "7. Data Model Overview")
    add_table(doc, ["Model", "Important fields", "Purpose"], model_rows(), widths=[1.25, 3.8, 2.45], font_size=7)

    add_heading(doc, "8. Route and API Map")
    add_table(doc, ["Method", "Endpoint", "Controller/middleware chain"], route_rows(), widths=[0.8, 2.7, 4.0], font_size=7)

    add_heading(doc, "9. Customer Frontend Routes")
    add_table(doc, ["Route", "Page", "How it works"], frontend_routes(), widths=[1.4, 1.8, 4.3], font_size=8)

    add_heading(doc, "10. Admin Dashboard Routes")
    add_table(doc, ["Route", "Page", "How it works"], admin_routes(), widths=[1.4, 1.8, 4.3], font_size=8)


def add_deep_flows(doc: Document) -> None:
    add_heading(doc, "11. Major Code Flows")

    flows = [
        (
            "Public Property Browsing",
            [
                "User opens /properties.",
                "PropertiesPage calls propertiesAPI.getAll(), which sends GET /api/products/list through frontend/src/services/api.ts.",
                "server.js has mounted productRoutes at /api/products, so /list reaches productController.listproperty.",
                "listproperty queries Property documents where status is active or legacy status is missing, sorts newest first, applies pagination, and returns { property, pagination, success }.",
                "PropertiesPage stores the results in state, useMemo filters by location/type/price/beds/baths/amenities, then PropertiesGrid/PropertyCard renders the UI.",
            ],
        ),
        (
            "Property Details and Appointment Booking",
            [
                "User clicks a property card and navigates to /property/:id.",
                "PropertyDetailsPage calls propertiesAPI.getById(id), which maps to GET /api/products/single/:id.",
                "productController.singleproperty finds the Property document and hides non-active pending/rejected/expired listings from public users.",
                "The page composes PropertyHeader, PropertyHeroImage, PropertyAbout, PropertyAmenities, PropertyLocation, and ScheduleViewingCard.",
                "ScheduleViewingCard posts booking data to /api/appointments/schedule; appointmentController.scheduleViewing creates an Appointment for guest or logged-in user and can trigger email notification logic.",
            ],
        ),
        (
            "Customer Authentication",
            [
                "SignUpPage and SignInPage render auth form components.",
                "AuthContext exposes register() and login(), both using userAPI from services/api.ts.",
                "Register posts { name, email, password } to /api/users/register; userController.register validates email, blocks disposable/fake emails, hashes the password with bcrypt, creates the user, and returns a JWT.",
                "Login posts credentials to /api/users/login; userController.login compares bcrypt hash and returns a JWT valid for 7 or 30 days.",
                "AuthContext saves brickbyte_token and brickbyte_user in localStorage, so apiClient attaches Authorization: Bearer <token> to future protected calls.",
            ],
        ),
        (
            "User Submitted Listing Approval Workflow",
            [
                "Authenticated user opens /add-property and uploads listing details/images.",
                "AddPropertyPage builds FormData and calls userListingsAPI.create().",
                "POST /api/user/properties runs protect middleware, then multer.array('images', 4), then propertyController.createUserListing.",
                "The controller uploads images to ImageKit, creates a Property with status='pending', postedBy=user._id, and expiresAt=45 days.",
                "Admin opens /pending-listings, calls /api/admin/properties/pending, then approves or rejects via PUT /api/admin/properties/:id/approve or /reject.",
                "Admin actions update the property status and write AdminActivityLog records for auditability.",
            ],
        ),
        (
            "AI Property Hub Search",
            [
                "AIPropertyHubPage only enables full functionality when VITE_ENABLE_AI_HUB=true.",
                "AIHeroSection collects city, locality, BHK, possession, budget, property type, and category. It requires localStorage API keys before submitting.",
                "frontend/src/services/api.ts posts /api/ai/search and attaches X-Github-Key and X-Firecrawl-Key headers from localStorage.",
                "propertyRoutes applies a distributed hourly limiter, transformAISearchRequest converts frontend price/type naming into backend fields, then propertyController.searchProperties runs.",
                "searchProperties requires both user-provided keys, builds a cache key, checks MongoDB SearchCache, and coalesces duplicate in-flight requests.",
                "FirecrawlService searches 99acres, MagicBricks, Housing.com, and optionally NoBroker, extracts structured listing JSON, filters rentals/PG/out-of-budget results, deduplicates, and interleaves sources.",
                "AIService sends cleaned property data to GitHub Models using GPT-4.1 mini with GPT-4.1 nano fallback and circuit breakers.",
                "validateAIResponse normalizes/repairs AI JSON so the frontend receives consistent properties + analysis.",
                "AISearchResults renders result cards and comparison UI; AIAnalysisPanel renders best value, recommendations, scores, red flags, and verdicts.",
            ],
        ),
        (
            "Location Trends Flow",
            [
                "After a successful AI search, the user can click Load Location Trends; it is intentionally not auto-fired.",
                "aiAPI.locationTrends(city) sends GET /api/locations/:city/trends with the same user API-key headers.",
                "Backend checks cache, runs FirecrawlService.getLocationTrends(), asks AIService.analyzeLocationTrends(), validates the response, caches it, and returns locations + analysis.",
                "AILocationTrends displays price per square foot, appreciation, rental yield, top appreciation, best rental yield, and tips.",
            ],
        ),
        (
            "Admin Dashboard Stats",
            [
                "Admin logs in at /login; login.jsx posts /api/users/admin and stores admin token + isAdmin flag.",
                "ProtectedRoute and AuthContext decode the JWT and block dashboard routes when token is missing/expired.",
                "Dashboard first calls enhanced endpoints /api/admin/stats/overview, /stats/users, /stats/properties, and /activity-logs?limit=10 using Promise.allSettled.",
                "If enhanced stats fail, Dashboard falls back to /api/admin/stats.",
                "Admin endpoints are protected by adminProtect, which verifies the JWT email matches ADMIN_EMAIL.",
            ],
        ),
    ]

    for title, steps in flows:
        add_heading(doc, title, level=2)
        add_numbered(doc, steps)

    add_heading(doc, "12. AI Hub Pipeline Diagram")
    add_code_block(
        doc,
        """
AI search form
  -> localStorage check for GitHub + Firecrawl keys
  -> POST /api/ai/search
  -> transformAISearchRequest
  -> distributed rate limiter
  -> SearchCache lookup
  -> requestCoalescer
  -> Firecrawl multi-source search/extraction
  -> filter/dedupe/interleave properties
  -> GitHub Models JSON analysis
  -> validate/fix AI response
  -> cache result
  -> React result cards + analysis sidebar
        """,
    )

    add_heading(doc, "13. Security, Reliability, and Performance")
    add_table(
        doc,
        ["Concern", "Where implemented", "Explanation"],
        [
            ["JWT auth", "userController.js, authMiddleware.js, AuthContext files", "Users and admins receive JWTs; frontend stores them in localStorage; Axios attaches Bearer tokens."],
            ["Admin protection", "adminRoutes.js + adminProtect", "All /api/admin routes require a token whose decoded email matches ADMIN_EMAIL."],
            ["CORS", "server.js", "Only known frontend/admin/local origins are allowed for browser requests."],
            ["Rate limiting", "server.js, rateLimitMiddleware.js, propertyRoutes.js, distributedRateLimiter.js", "General requests and expensive AI requests are capped to protect API cost and server capacity."],
            ["NoSQL injection mitigation", "express-mongo-sanitize in server.js", "Strips/replaces malicious MongoDB operator keys from user input."],
            ["Request tracing", "requestIdMiddleware.js + logger.js", "Each request gets an ID for logs and error responses."],
            ["Structured logging", "utils/logger.js", "Winston logger records JSON-like operational data and errors."],
            ["Caching", "searchCacheModel.js + propertyController.js", "AI search and trends are cached in MongoDB with TTL to save credits and improve speed."],
            ["Request coalescing", "utils/requestCoalescer.js", "Identical concurrent AI requests share the same in-flight promise."],
            ["Circuit breakers", "utils/circuitBreaker.js, aiService.js, firecrawlService.js", "Stops repeatedly calling failing third-party services until a timeout passes."],
            ["Cron cleanup", "expireListings.js, autoUnsuspend.js", "Background jobs expire old listings and automatically unsuspend users after suspension end date."],
            ["Image uploads", "multer.js + imagekit.js", "Multer receives temporary files; ImageKit stores/CDN-serves the uploaded property images."],
        ],
        widths=[1.5, 2.5, 3.5],
        font_size=8,
    )

    add_heading(doc, "14. Things to Watch or Fix")
    add_bullets(
        doc,
        [
            "backend/middleware/authMiddleware.js defines checkAppointmentOwnership but references Appointment without importing it; that function is not currently used by routes, but it would fail if wired in.",
            "frontend/src/services/api.ts passes rememberMe in login calls from AuthContext, but the TypeScript signature only lists email/password; TypeScript may complain depending on compiler strictness.",
            "AI Hub and API comments contain some mojibake/encoding artifacts in comments and labels; functionality is mostly unaffected, but clean text is better for a final submission.",
            "The repo contains duplicated source under Real-Estate-Website-main, logs, build output, local MongoDB downloads, and node_modules. Clean these before submitting or deploying your own copy.",
            "No real automated test suite is present; package scripts say no tests specified. For a presentation, be ready to explain manual testing and health-check validation.",
        ],
    )


def add_file_inventory(doc: Document, files: list[Path]) -> None:
    add_page_break(doc)
    add_heading(doc, "15. File-by-File Guide")
    add_para(
        doc,
        "This appendix lists every active project-owned file I analyzed, excluding vendor dependencies, build output, logs, local MongoDB binary packages, and the duplicated Real-Estate-Website-main snapshot. The duplicate snapshot mirrors the same folders/files and should be explained as a copied source backup."
    )

    groups = summarize_inventory(files)
    order = [
        "repo root docs",
        "github workflow docs",
        "backend root/deploy",
        "backend config",
        "backend routes",
        "backend controllers",
        "backend models",
        "backend services",
        "backend middleware",
        "backend utils",
        "frontend root/deploy",
        "frontend source core",
        "frontend pages",
        "frontend components",
        "frontend ui primitives",
        "frontend public assets",
        "admin root/deploy",
        "admin source core",
        "admin pages",
        "admin components",
        "shared utilities",
        "documentation images",
    ]
    for group in order:
        paths = groups.get(group, [])
        if not paths:
            continue
        add_heading(doc, group.replace("/", " ").title(), level=2)
        rows = []
        for p in paths:
            rows.append([
                rel(p),
                classify_language(p),
                extract_symbols(p) or extract_api_calls(p) or "",
                purpose_for(p),
            ])
        add_table(doc, ["File", "Type", "Key symbols/calls", "Purpose"], rows, widths=[2.35, 1.15, 1.7, 3.0], font_size=6)


def add_presentation_section(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "16. Presentation Preparation")
    add_para(
        doc,
        "The best way to present this project is to explain it as three apps connected by one backend: customer frontend, admin dashboard, and API/server. Then show one complete feature flow, such as AI search or listing approval."
    )

    add_heading(doc, "Important Topics You Could Be Asked")
    topics = [
        "Why did you choose React and Vite for the frontend?",
        "What is the difference between the customer frontend and the admin dashboard?",
        "How does React Router map URLs to pages in App.tsx?",
        "How does index.html connect to main.tsx and App.tsx?",
        "What is the role of AuthContext, and why is localStorage used?",
        "How is JWT authentication implemented?",
        "How does the backend protect admin-only routes?",
        "What are Express middleware, routes, controllers, models, services, and utils?",
        "How does Mongoose connect JavaScript objects to MongoDB collections?",
        "What fields exist in the Property model, and how does status control visibility?",
        "How does a user-submitted listing become visible publicly?",
        "How are images uploaded and where are they stored?",
        "What is ImageKit and why use a CDN?",
        "How does appointment booking work for guest and logged-in users?",
        "What is the AI Hub feature flag VITE_ENABLE_AI_HUB?",
        "Why does the AI Hub require user-provided Firecrawl and GitHub Models keys?",
        "How does Firecrawl collect property data from portals?",
        "How does the backend filter rentals/PG listings and out-of-budget results?",
        "How does the AI service rank properties and return JSON analysis?",
        "What is caching and why is MongoDB TTL useful for AI search?",
        "What is request coalescing and how does it save API calls?",
        "What is a circuit breaker and why is it used for third-party APIs?",
        "How does rate limiting protect the system?",
        "What is CORS and why must frontend/admin origins be allowlisted?",
        "What is Helmet and what security problem does it address?",
        "What are the deployment targets: Vercel for frontend/admin and Render for backend?",
        "What environment variables are required locally?",
        "What would you improve if you had more time?",
        "How would you add tests to this project?",
        "How would you clean the repository before submitting it as your own maintained version?",
    ]
    add_bullets(doc, topics)

    add_heading(doc, "Short Presentation Script")
    add_numbered(
        doc,
        [
            "Start with the problem: users need searchable, structured, trustworthy real estate listings and admins need a way to manage listings and users.",
            "Show the architecture: customer React app and admin React app both call the Express backend, which stores data in MongoDB and integrates with ImageKit, email, Firecrawl, and AI.",
            "Demo browsing: /properties fetches active listings, filters them on the frontend, and links to /property/:id.",
            "Demo user workflow: signup/login, submit a property, status becomes pending, admin approves it, then public listing appears.",
            "Demo AI workflow: user enters city/budget, frontend sends API keys in headers, backend scrapes listings, AI ranks them, cache stores the result.",
            "Close with security/reliability: JWT, admin guards, CORS, Helmet, rate limits, Mongo sanitize, logging, health checks, caching, and circuit breakers.",
        ],
    )

    add_heading(doc, "One-Minute Architecture Answer")
    add_para(
        doc,
        "BrickByte is a monorepo with a customer React/TypeScript frontend, a separate React/JavaScript admin dashboard, and a Node/Express backend. The frontend and admin are single-page applications built by Vite. Both communicate with the backend through Axios. The backend mounts route modules under /api, uses controllers for business logic, Mongoose models for MongoDB data, and services for external integrations such as ImageKit, email, Firecrawl, and GitHub Models. Authentication is JWT-based; user tokens protect listing and appointment features, while admin tokens protect /api/admin routes. The AI Hub uses user-supplied API keys, scrapes property portals through Firecrawl, ranks results with GPT-4.1, caches responses in MongoDB, and displays analysis in React."
    )


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    files = iter_project_files()

    doc = Document()
    configure_doc(doc)

    add_cover(doc, len(files))
    add_page_break(doc)
    add_architecture_sections(doc, files)
    add_deep_flows(doc)
    add_file_inventory(doc, files)
    add_presentation_section(doc)

    doc.save(OUT_PATH)
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
